"""
Report Generator — Creates a professional DOCX report matching the
"Investment Proposal" reference design (Go.pdf).

Design system:
  • Deep navy hero banner with white text
  • 4-metric summary strip below the banner
  • ALL-CAPS section headers with a navy underline rule
  • Tables: navy header row, zebra-striped body, no vertical borders
  • Multi-year comparison columns (FY2024 | FY2025)
  • Growth metrics callout
  • Ratio color coding: green = PASS, red = FAIL
  • Risk cards: red left-border; Positive cards: green left-border
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from ratios import ratios_to_flat_list


# ── Color Palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x36, 0x5D)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0xF7, 0xFA, 0xFC)
GREEN      = RGBColor(0x38, 0xA1, 0x69)
RED        = RGBColor(0xE5, 0x3E, 0x3E)
ORANGE     = RGBColor(0xED, 0x89, 0x36)
MUTED      = RGBColor(0x71, 0x80, 0x96)

NAVY_HEX   = "1A365D"
LGRAY_HEX  = "F7FAFC"
GREEN_HEX  = "E6F4EA"
RED_HEX    = "FDE8E8"


# ── Low-level Helpers ────────────────────────────────────────────────────────

def _shade(cell, hex_color: str):
    tc_pr = cell._element.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tc_pr.append(shading)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc_pr = cell._element.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(margins)


def _remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    for old in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(old)
    tbl_pr.append(borders)


def _set_table_horizontal_borders_only(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D5DD"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    for old in tbl_pr.findall(qn('w:tblBorders')):
        tbl_pr.remove(old)
    tbl_pr.append(borders)


def _run(para, text, size=11, bold=False, italic=False, color=BLACK, font="Calibri"):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def _section_heading(doc, title: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(2)
    r = para.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(12)
    pPr = rule._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{NAVY_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return para


def _data_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_horizontal_borders_only(table)

    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _shade(cell, NAVY_HEX)
        _set_cell_margins(cell, 60, 60, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, hdr, size=9, bold=True, color=WHITE)

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            _set_cell_margins(cell, 50, 50, 100, 100)
            if ri % 2 == 1:
                _shade(cell, LGRAY_HEX)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            val_str = str(val) if val else "—"
            if "PASS" in val_str.upper():
                _run(p, val_str, size=10, bold=True, color=GREEN)
            elif "FAIL" in val_str.upper():
                _run(p, val_str, size=10, bold=True, color=RED)
            elif "CAUTION" in val_str.upper():
                _run(p, val_str, size=10, bold=True, color=ORANGE)
            else:
                _run(p, val_str, size=10, color=BLACK)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def _fmt_currency(val, currency="₹"):
    try:
        v = float(val)
        if v == 0:
            return "—"
        if abs(v) >= 1_00_00_000:
            return f"{currency}{v / 1_00_00_000:,.2f}Cr"
        elif abs(v) >= 1_00_000:
            return f"{currency}{v / 1_00_000:,.2f}L"
        else:
            return f"{currency}{v:,.2f}"
    except (ValueError, TypeError):
        return str(val) if val else "—"


def _add_body(doc, text: str):
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    _run(p, text, size=11, color=BLACK)


def _observations_text(section_data: dict) -> str:
    if not section_data:
        return "Analysis not available."
    overall = section_data.get("overall", "")
    obs = section_data.get("observations", [])
    parts = []
    if overall:
        parts.append(f"Overall assessment: {overall}.")
    if obs:
        parts.append(" ".join(obs))
    return " ".join(parts) if parts else "Analysis not available."


def _get_ratio_value_multi(computed_ratios, year, category, name):
    yr_data = computed_ratios.get(year, {})
    cat = yr_data.get(category, {})
    ratio = cat.get(name, {})
    return ratio.get("formatted", "N/A")


# ── Main Generator ───────────────────────────────────────────────────────────

def generate_report(analysis: dict, output_dir: str = ".") -> str:
    doc = Document()

    company = analysis.get("company_name", "Unknown Company")
    date_str = datetime.now().strftime("%d %b %Y")
    safe_company = company.replace("/", "-").replace("\\", "-")[:50]
    filename = f"{safe_company}_Financial_Analysis_{datetime.now().strftime('%Y%m%d')}.docx"

    # Global style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.3
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)

    rec = analysis.get("recommendation", {})
    financials = analysis.get("financials", {})
    bg = analysis.get("company_background", {})
    fin_analysis = analysis.get("financial_analysis", {})
    risk = analysis.get("risk_analysis", {})
    comp_data = analysis.get("competitor_analysis", {})
    computed_ratios = analysis.get("computed_ratios", {})
    growth = analysis.get("growth_metrics", {})
    proj = analysis.get("projection_analysis", {})

    csym = "₹"
    verdict = rec.get("recommendation", "N/A")

    # ── Multi-year setup ─────────────────────────────────────────────────────
    years = financials.get("years_found", [])
    sorted_years = sorted(years) if years else ["FY2025"]
    latest_year = sorted_years[-1]
    latest = financials.get(latest_year, {})
    if not isinstance(latest, dict):
        latest = financials  # legacy fallback

    def _fy_val(field, year=None):
        yr = year or latest_year
        d = financials.get(yr, {})
        return d.get(field, 0) if isinstance(d, dict) else 0

    def _multi_rows(fields):
        """Build multi-year table rows: [("Label","key"), ...]"""
        hdrs = ["Particulars"] + sorted_years
        rows = []
        for label, key in fields:
            row = [label]
            for yr in sorted_years:
                d = financials.get(yr, {})
                row.append(_fmt_currency(d.get(key, 0), csym) if isinstance(d, dict) else "—")
            rows.append(row)
        return hdrs, rows

    net_worth = _fy_val("equity")
    col_w = [8] + [4.5] * len(sorted_years)

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. HERO BANNER
    # ═══════════════════════════════════════════════════════════════════════════
    banner = doc.add_table(rows=1, cols=2)
    _remove_table_borders(banner)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER

    left = banner.rows[0].cells[0]
    left.text = ""
    _shade(left, NAVY_HEX)
    _set_cell_margins(left, 200, 200, 200, 100)
    left.width = Cm(12)

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    _run(p, "REHBAR FINANCIAL SERVICES — CONFIDENTIAL", size=8, color=RGBColor(0xA0, 0xAE, 0xC0))
    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    _run(p2, "INVESTMENT PROPOSAL", size=22, bold=True, color=WHITE)
    p3 = left.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    _run(p3, company, size=13, bold=True, color=RGBColor(0xB0, 0xC4, 0xDE))

    desc = bg.get("company_description", "")
    if desc:
        pd = left.add_paragraph()
        pd.paragraph_format.space_after = Pt(4)
        _run(pd, desc[:250] + ("..." if len(desc) > 250 else ""), size=9, color=RGBColor(0xCC, 0xD5, 0xE0))

    right = banner.rows[0].cells[1]
    right.text = ""
    _shade(right, NAVY_HEX)
    _set_cell_margins(right, 200, 200, 100, 200)
    right.width = Cm(5)

    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_after = Pt(2)
    _run(pr, f"Generated: {date_str}", size=8, color=RGBColor(0xA0, 0xAE, 0xC0))

    pbadge = right.add_paragraph()
    pbadge.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pbadge.paragraph_format.space_before = Pt(8)
    pbadge.paragraph_format.space_after = Pt(8)
    badge_color = GREEN if verdict.upper() in ("BUY", "CONDITIONAL APPROVAL") else RED if verdict.upper() in ("SELL", "AVOID") else ORANGE
    _run(pbadge, f" {verdict.upper()} ", size=13, bold=True, color=badge_color)

    pexp = right.add_paragraph()
    pexp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(pexp, "Net Worth", size=8, color=RGBColor(0xA0, 0xAE, 0xC0))
    pexp2 = right.add_paragraph()
    pexp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(pexp2, _fmt_currency(net_worth, csym), size=16, bold=True, color=WHITE)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. KEY METRICS STRIP
    # ═══════════════════════════════════════════════════════════════════════════
    cur_r = _get_ratio_value_multi(computed_ratios, latest_year, "Liquidity Ratios", "Current Ratio")
    rev_g = growth.get("revenue_growth", "—")
    profit_g = growth.get("net_profit_growth", "—")

    metrics_data = [
        (f"NET WORTH ({latest_year})", _fmt_currency(net_worth, csym)),
        (f"EBITDA ({latest_year})", _fmt_currency(_fy_val("ebitda"), csym)),
        ("REVENUE GROWTH", rev_g if rev_g != "—" else "N/A"),
        (f"CURRENT RATIO ({latest_year})", cur_r),
    ]

    metrics = doc.add_table(rows=2, cols=4)
    _remove_table_borders(metrics)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(metrics_data):
        lc = metrics.rows[0].cells[i]
        lc.text = ""
        _set_cell_margins(lc, 60, 0, 80, 80)
        _run(lc.paragraphs[0], label, size=8, bold=True, color=MUTED)
        vc = metrics.rows[1].cells[i]
        vc.text = ""
        _set_cell_margins(vc, 0, 80, 80, 80)
        _run(vc.paragraphs[0], value, size=16, bold=True, color=NAVY)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. COMPANY BACKGROUND
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Company Background")
    _add_body(doc, bg.get("company_description", "Company background not available."))

    products = bg.get("key_products_services", [])
    if products:
        p = doc.add_paragraph()
        _run(p, "Key Products & Services: ", size=11, bold=True, color=NAVY)
        _run(p, ", ".join(products), size=11)

    mgmt = bg.get("key_management", [])
    if mgmt:
        p = doc.add_paragraph()
        _run(p, "Key Management: ", size=11, bold=True, color=NAVY)
        _run(p, ", ".join(mgmt), size=11)

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. SECTOR & MARKET CONTEXT
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Sector & Market Context")
    _add_body(doc, comp_data.get("industry_overview", "") or bg.get("market_position", "") or "Sector context not available.")

    if comp_data.get("competitors"):
        _data_table(
            doc,
            ["Competitor", "Description", "Market Position"],
            [[c.get("name", ""), c.get("description", "")[:80], c.get("market_position", "")] for c in comp_data["competitors"]],
            [4, 9, 4],
        )

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. FINANCIAL OVERVIEW (multi-year)
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Financial Overview")
    _add_body(doc, fin_analysis.get("executive_summary", "Financial overview not available."))

    # Growth callout
    if growth and growth.get("revenue_growth"):
        gp = doc.add_paragraph()
        gp.paragraph_format.space_after = Pt(8)
        _run(gp, f"YoY Growth ({growth.get('comparison', '')}): ", size=11, bold=True, color=NAVY)
        _run(gp, f"Revenue {growth.get('revenue_growth', 'N/A')}", size=11, bold=True, color=GREEN)
        _run(gp, "  |  ", size=11, color=MUTED)
        _run(gp, f"Net Profit {growth.get('net_profit_growth', 'N/A')}", size=11, bold=True, color=GREEN)
        _run(gp, "  |  ", size=11, color=MUTED)
        _run(gp, f"Equity {growth.get('equity_growth', 'N/A')}", size=11, bold=True, color=NAVY)

    headers, rows = _multi_rows([
        ("Net Worth (Equity)", "equity"),
        ("Total Revenue", "revenue"),
        ("Cost of Materials", "cost_of_materials"),
        ("EBITDA", "ebitda"),
        ("Net Profit", "net_profit"),
    ])
    _data_table(doc, headers, rows, col_w)

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. PROFITABILITY ANALYSIS (multi-year)
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Profitability Analysis")
    _add_body(doc, _observations_text(fin_analysis.get("profitability_assessment", {})))

    headers, rows = _multi_rows([
        ("Revenue", "revenue"),
        ("Total Expenses", "total_expenses"),
        ("Finance Costs", "finance_cost"),
        ("Depreciation", "depreciation"),
        ("EBITDA", "ebitda"),
        ("Net Profit", "net_profit"),
    ])
    _data_table(doc, headers, rows, col_w)

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. BALANCE SHEET ANALYSIS (multi-year)
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Balance Sheet Analysis")
    _add_body(doc, _observations_text(fin_analysis.get("balance_sheet_strength", {})))

    headers, rows = _multi_rows([
        ("Net Worth (Equity)", "equity"),
        ("Total Debt", "total_debt"),
        ("Total Assets", "total_assets"),
        ("Current Assets", "current_assets_total"),
        ("Current Liabilities", "current_liabilities_total"),
        ("Tangible Assets", "tangible_assets"),
    ])
    _data_table(doc, headers, rows, col_w)

    # ═══════════════════════════════════════════════════════════════════════════
    # 8. CASH FLOW ANALYSIS (multi-year)
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Cash Flow Analysis")
    _add_body(doc, _observations_text(fin_analysis.get("cash_flow_assessment", {})))

    headers, rows = _multi_rows([
        ("Cash & Equivalents", "cash_and_equivalents"),
        ("Trade Receivables", "trade_receivables"),
        ("Inventories", "inventories"),
    ])
    _data_table(doc, headers, rows, col_w)

    # ═══════════════════════════════════════════════════════════════════════════
    # 9. KEY FINANCIAL RATIOS (multi-year, color-coded)
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _section_heading(doc, "Key Financial Ratios")
    _add_body(doc, "Values in green indicate PASS (within benchmark), red indicates FAIL.")

    legend = doc.add_paragraph()
    legend.paragraph_format.space_after = Pt(8)
    _run(legend, "✅ PASS  |  ❌ FAIL", size=9, italic=True, color=MUTED)

    # Build multiple multi-year ratio tables (one per category)
    ratio_headers = ["Ratio"] + sorted_years
    latest_ratios = computed_ratios.get(latest_year, {})
    for category, items in latest_ratios.items():
        # Add category subheader
        ch = doc.add_paragraph()
        ch.paragraph_format.space_before = Pt(16)
        ch.paragraph_format.space_after = Pt(4)
        _run(ch, category, size=12, bold=True, color=NAVY)

        ratio_rows = []
        for name, data in items.items():
            row = [name]
            for yr in sorted_years:
                yr_ratios = computed_ratios.get(yr, {})
                yr_cat = yr_ratios.get(category, {})
                yr_data = yr_cat.get(name, {})
                val = yr_data.get("formatted", "N/A")
                status = yr_data.get("status", "")
                if "PASS" in status:
                    row.append(f"✅ {val}")
                elif "FAIL" in status:
                    row.append(f"❌ {val}")
                else:
                    row.append(val)
            ratio_rows.append(row)

        _data_table(doc, ratio_headers, ratio_rows, col_w)

    # ═══════════════════════════════════════════════════════════════════════════
    # 10. RISK FACTORS
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Risk Factors")

    risk_factors = risk.get("risk_factors", [])
    if risk_factors:
        for i, rf in enumerate(risk_factors):
            severity = rf.get("severity", "Medium").upper()
            desc_text = rf.get("description", "")
            card = doc.add_table(rows=1, cols=1)
            _remove_table_borders(card)
            cell = card.rows[0].cells[0]
            cell.text = ""
            _shade(cell, RED_HEX)
            _set_cell_margins(cell, 80, 80, 120, 120)
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:start w:val="single" w:sz="24" w:space="0" w:color="E53E3E"/>'
                f'</w:tcBorders>'
            )
            tc_pr.append(tc_borders)
            p = cell.paragraphs[0]
            if severity == "HIGH":
                _run(p, f" {severity} ", size=8, bold=True, color=RED)
                _run(p, "  ", size=8)
            _run(p, f"{i + 1}. {desc_text}", size=10, color=BLACK)
            doc.add_paragraph()
    else:
        _add_body(doc, "No significant risk factors identified.")

    # ═══════════════════════════════════════════════════════════════════════════
    # 11. POSITIVE INDICATORS
    # ═══════════════════════════════════════════════════════════════════════════
    _section_heading(doc, "Positive Indicators")

    pros = risk.get("investment_pros", [])
    if pros:
        for i, pro in enumerate(pros):
            desc_text = pro.get("description", "")
            card = doc.add_table(rows=1, cols=1)
            _remove_table_borders(card)
            cell = card.rows[0].cells[0]
            cell.text = ""
            _shade(cell, GREEN_HEX)
            _set_cell_margins(cell, 80, 80, 120, 120)
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:start w:val="single" w:sz="24" w:space="0" w:color="38A169"/>'
                f'</w:tcBorders>'
            )
            tc_pr.append(tc_borders)
            p = cell.paragraphs[0]
            _run(p, "+  ", size=11, bold=True, color=GREEN)
            _run(p, f"{i + 1}. {desc_text}", size=10, color=BLACK)
            doc.add_paragraph()
    else:
        _add_body(doc, "No positive indicators identified.")

    # ═══════════════════════════════════════════════════════════════════════════
    # 12. INVESTMENT RECOMMENDATION
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _section_heading(doc, "Investment Recommendation")

    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vp.paragraph_format.space_before = Pt(12)
    vp.paragraph_format.space_after = Pt(4)
    v_color = GREEN if verdict.upper() in ("BUY", "CONDITIONAL APPROVAL") else RED if verdict.upper() in ("SELL", "AVOID") else ORANGE
    _run(vp, verdict.upper(), size=36, bold=True, color=v_color)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    _run(meta, f"Confidence: {rec.get('confidence_level', 'N/A')}  •  ", size=11, color=MUTED)
    _run(meta, f"Horizon: {rec.get('target_horizon', 'N/A')}  •  ", size=11, color=MUTED)
    _run(meta, f"Suitable for: {rec.get('suitable_for', 'N/A')}", size=11, color=MUTED)

    if rec.get("summary"):
        _add_body(doc, rec["summary"])
    if rec.get("detailed_rationale"):
        doc.add_paragraph()
        _add_body(doc, rec["detailed_rationale"])

    reasons = rec.get("key_reasons", [])
    if reasons:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        _run(p, "Key Reasons:", size=12, bold=True, color=NAVY)
        for r in reasons:
            doc.add_paragraph(r, style="List Bullet")

    caveats = rec.get("caveats", [])
    if caveats:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        _run(p, "Caveats:", size=12, bold=True, color=ORANGE)
        for c in caveats:
            doc.add_paragraph(c, style="List Bullet")

    # ═══════════════════════════════════════════════════════════════════════════
    # 13. PROJECTION ANALYSIS
    # ═══════════════════════════════════════════════════════════════════════════
    if proj and proj.get("review_table"):
        doc.add_page_break()
        _section_heading(doc, "Management Projection Analysis")

        # Credibility badge line
        cred = proj.get("overall_credibility", "N/A")
        period = proj.get("projection_period", "")
        cred_line = doc.add_paragraph()
        cred_color = GREEN if cred == "Realistic" else RED if cred == "Optimistic" else ORANGE
        cred_line.paragraph_format.space_after = Pt(6)
        _run(cred_line, f"Overall Credibility: ", size=11, bold=True, color=NAVY)
        _run(cred_line, cred, size=11, bold=True, color=cred_color)
        if period:
            _run(cred_line, f"  •  Period: {period}", size=11, color=MUTED)

        if proj.get("overall_credibility_summary"):
            _add_body(doc, proj["overall_credibility_summary"])

        # Management assumptions
        assumptions = proj.get("management_assumptions", [])
        if assumptions:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            _run(p, "Management Assumptions:", size=12, bold=True, color=NAVY)
            for a in assumptions:
                doc.add_paragraph(a, style="List Bullet")

        # Projection review table
        review_table = proj.get("review_table", [])
        if review_table:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            _run(p, "Projection Review:", size=12, bold=True, color=NAVY)
            headers = ["Metric", "Management Projection", "Historical Baseline", "Credibility", "Rationale"]
            rows = []
            for row in review_table:
                flag = " ⚠" if row.get("risk_flag") else ""
                cred_val = row.get("credibility", "—")
                rows.append([
                    row.get("metric", "—") + flag,
                    row.get("management_projection", "—"),
                    row.get("historical_baseline", "—"),
                    cred_val,
                    (row.get("credibility_reason", "—") or "—")[:120],
                ])
            _data_table(doc, headers, rows, [4, 3, 3, 2.5, 5])

        # Key concerns
        concerns = proj.get("key_concerns", [])
        if concerns:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            _run(p, "Key Concerns:", size=12, bold=True, color=RED)
            for c in concerns:
                doc.add_paragraph(f"⚠  {c}", style="List Bullet")

        # AI Counter-Projection
        cp = proj.get("ai_counter_projection", {})
        if cp:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            _run(p, "AI Counter-Projection", size=13, bold=True, color=NAVY)
            if cp.get("methodology"):
                _add_body(doc, cp["methodology"])
            cp_rows = cp.get("projections", [])
            if cp_rows:
                # Build a year-keyed table
                years_cp = cp_rows[0].get("year_by_year", []) if cp_rows else []
                yr_labels = [y.get("year", "") for y in years_cp]
                cp_headers = ["Metric"] + yr_labels
                cp_table_rows = []
                for item in cp_rows:
                    row = [item.get("metric", "—")]
                    for y in item.get("year_by_year", []):
                        row.append(y.get("value", "—"))
                    cp_table_rows.append(row)
                if cp_table_rows:
                    col_w_cp = [5.5] + [3] * len(yr_labels)
                    _data_table(doc, cp_headers, cp_table_rows, col_w_cp)
            if cp.get("summary"):
                _add_body(doc, cp["summary"])

    # ═══════════════════════════════════════════════════════════════════════════
    # 14. DISCLAIMER
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    _section_heading(doc, "Disclaimer")
    disclaimer = (
        "This report has been generated using AI-powered analysis tools and is intended for "
        "informational purposes only. It does not constitute financial advice, investment guidance, "
        "or a recommendation to buy, sell, or hold any securities. "
        "Investors should conduct their own independent research and consult with qualified financial "
        "advisors before making any investment decisions. Past financial performance does not guarantee "
        "future results.\n\n"
        f"Report generated on {date_str} using automated financial analysis tools."
    )
    p = doc.add_paragraph()
    _run(p, disclaimer, size=9, italic=True, color=MUTED)

    # ── Save ─────────────────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"\n📄 Report saved: {output_path}")
    return output_path
